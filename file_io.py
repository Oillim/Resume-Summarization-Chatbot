import os
import PyPDF2
import docx
import openai
from config import OPENAI_API_KEY
from database import Database
from model import Model
class FileHandler:
    def __init__(self, resume_folder="resumes"):
        self.resume_folder = resume_folder
        os.makedirs(self.resume_folder, exist_ok=True)
        openai.api_key = OPENAI_API_KEY
        self.db = Database()
        self.model = Model()

    def save_file(self, file, file_name):
        """Save the uploaded file to the local folder"""
        file_path = os.path.join(self.resume_folder, file_name)
        with open(file_path, "wb") as f:
            f.write(file.read())
        return file_path

    def extract_text_from_pdf(self, pdf_path):
        """Extract text from a PDF file"""
        with open(pdf_path, "rb") as file:
            reader = PyPDF2.PdfReader(file)
            text = ""
            for page in reader.pages:
                text += page.extract_text()
        return text

    def extract_text_from_docx(self, docx_path):
        """Extract text from a DOCX file"""
        doc = docx.Document(docx_path)
        return "\n".join([para.text for para in doc.paragraphs])

    def extract_text(self, file_path):
        """Extract text from PDF or DOCX files"""
        if file_path.endswith(".pdf"):
            return self.extract_text_from_pdf(file_path)
        elif file_path.endswith(".docx"):
            return self.extract_text_from_docx(file_path)
        else:
            raise ValueError("Unsupported file format")
    
    def read_file(self, file):
        file_extension = file.name.split(".")[-1]
        if file_extension == "pdf":
            reader = PyPDF2.PdfReader(file)
            content = ""
            for page in reader.pages:
                content += page.extract_text()
            return content
        elif file_extension == "docx":
            doc = docx.Document(file)
            return "\n".join([para.text for para in doc.paragraphs])
        else:
            return "Unsupported file type"
        
    def store_in_database(self, data):
        """Store the extracted data into the MySQL database"""
        query = """
            INSERT INTO resumes (name, email, phone, skills, experience, education, summary)
            VALUES (%s, %s, %s, %s, %s, %s, %s)

        """
        params = (
            data.get("name", ""),
            data.get("email", ""),
            data.get("phone", ""),
            data.get("skills", ""),
            data.get("experience", ""),
            data.get("education", ""),
            data.get("summary", ""),
            
        )
        self.db.execute_query(query, params)

    def process_resume(self, file):
        """Process the resume file: Extract text, use OpenAI API, and store in DB"""
        file_path = self.save_file(file, file.name)
        extracted_text = self.extract_text(file_path)
        
        # Process the extracted text using OpenAI to get structured data
        extracted_data = self.model.process_with_openai(extracted_text)
        self.store_in_database(extracted_data)
        return extracted_data
